import io
import os
from datetime import datetime, timezone
from typing import Optional

from fastapi import Depends, FastAPI, File, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

from app.pipeline import (
    InputValidationError,
    PipelineOutput,
    ScoringError,
    generate_career_dna_report,
    serialise_output,
    serialise_pipeline_error,
)
from app.schemas import RawInput
from app.models.client import (
    ClientRecord, ClientProfile, CreateClientRequest, UpdateClientRequest,
    MarketRadarRequest, Opportunity, OpportunityRequest,
    SessionNote, ActionItem, SessionNoteRequest, ActionItemRequest,
    AdvisorBrief, TargetContact, TargetContactRequest, ContactSearchRequest,
    ExtractFromNotesRequest,
)
from app.data.storage_selector import list_clients, get_client, create_client, update_client, delete_client
from app.config import get_storage_backend, StorageConfigError
from app.data.storage import check_connectivity as _json_storage_connectivity
from app.db.session import check_connectivity as _postgres_storage_connectivity
from app.services.positioning import generate_positioning
from app.services.cv_intelligence import analyse_cv
from app.services.market_radar import run_market_radar
from app.services.advisor_brief import generate_advisor_brief
from app.services.contact_search import search_contacts
from app.services.notes_extraction import extract_network_from_notes

app = FastAPI(title="Career DNA API")


def verify_trial_key(x_trial_key: Optional[str] = Header(default=None)) -> None:
    """
    Validates the X-Trial-Key header against the TRIAL_API_KEY env var.
    If TRIAL_API_KEY is not set (local dev), all requests are allowed.
    If it is set (production), a missing or wrong key returns 401.
    """
    trial_key = os.getenv("TRIAL_API_KEY")
    if not trial_key:
        return  # dev mode — no key configured, open access
    if x_trial_key != trial_key:
        raise HTTPException(status_code=401, detail="Unauthorized")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "https://career-dna-production.up.railway.app",
        "*",
    ],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def index():
    return {"status": "running"}


@app.get("/health")
def health():
    """Minimal health probe, safe to expose publicly (e.g. Railway health
    checks). Never returns DATABASE_URL, credentials, file paths, client
    counts, or any client content — only status, the configured storage
    backend name, and whether that backend is currently reachable."""
    try:
        backend = get_storage_backend()
    except StorageConfigError:
        return JSONResponse(
            status_code=503,
            content={"status": "error", "storage_backend": "unconfigured", "storage_connected": False},
        )

    try:
        connected = (
            _postgres_storage_connectivity() if backend == "postgres" else _json_storage_connectivity()
        )
    except Exception:
        connected = False

    return JSONResponse(
        status_code=200 if connected else 503,
        content={
            "status": "ok" if connected else "error",
            "storage_backend": backend,
            "storage_connected": connected,
        },
    )


@app.post("/generate-dna", response_model=None, dependencies=[Depends(verify_trial_key)])
def generate_dna(input_data: RawInput):
    try:
        output: PipelineOutput = generate_career_dna_report(input_data)
        return JSONResponse(status_code=200, content=serialise_output(output))

    except InputValidationError as exc:
        return JSONResponse(
            status_code=422,
            content=serialise_pipeline_error(exc),
        )

    except ScoringError as exc:
        return JSONResponse(
            status_code=422,
            content=serialise_pipeline_error(exc),
        )

    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={
                "error": True,
                "stage": "unknown",
                "message": f"Unexpected error: {exc}",
                "recoverable": False,
            },
        )


# ── Client workspace endpoints ────────────────────────────────────────────────


@app.get("/clients")
def get_clients():
    clients = list_clients()
    return JSONResponse(
        status_code=200,
        content=[c.model_dump(mode="json") for c in clients],
    )


@app.post("/clients")
def post_client(body: CreateClientRequest):
    if not body.name.strip():
        raise HTTPException(status_code=422, detail="Client name is required.")
    profile = ClientProfile(name=body.name.strip())
    record = ClientRecord(profile=profile)
    created = create_client(record)
    return JSONResponse(status_code=201, content=created.model_dump(mode="json"))


@app.get("/clients/{client_id}")
def get_client_by_id(client_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    return JSONResponse(status_code=200, content=record.model_dump(mode="json"))


@app.put("/clients/{client_id}")
def put_client(client_id: str, body: UpdateClientRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    record.profile = body.profile
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.delete("/clients/{client_id}")
def delete_client_by_id(client_id: str):
    if not delete_client(client_id):
        raise HTTPException(status_code=404, detail="Client not found.")
    return JSONResponse(status_code=200, content={"deleted": True})


_CV_FILE_MAX_BYTES = 10 * 1024 * 1024   # 10 MB
_CV_TEXT_MIN_CHARS = 200                 # below this we warn about scanned PDFs


@app.post("/clients/{client_id}/cv/extract-file")
async def post_extract_cv_file(client_id: str, file: UploadFile = File(...)):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")

    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ("pdf", "docx", "txt"):
        raise HTTPException(
            status_code=422,
            detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file.",
        )

    content = await file.read()
    if len(content) > _CV_FILE_MAX_BYTES:
        raise HTTPException(
            status_code=422,
            detail="File too large. Maximum size is 10 MB.",
        )

    text = ""
    warning = ""

    if ext == "txt":
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1", errors="replace")

    elif ext == "pdf":
        try:
            import pypdf  # noqa: PLC0415
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF extraction library not available.")
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not read PDF: {exc}")
        if len(text.strip()) < _CV_TEXT_MIN_CHARS:
            warning = (
                "We couldn't extract enough text from this PDF. "
                "It may be scanned or image-based. "
                "Please paste the CV text manually."
            )

    elif ext == "docx":
        try:
            import docx  # noqa: PLC0415
            doc = docx.Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts).strip()
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX extraction library not available.")
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"Could not read DOCX file: {exc}")
        if len(text.strip()) < 100:
            warning = (
                "Very little text was extracted from this document. "
                "Please check the file or paste the CV text manually."
            )

    return JSONResponse(status_code=200, content={"text": text, "warning": warning})


@app.post("/clients/{client_id}/analyse-cv")
def post_analyse_cv(client_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    try:
        result = analyse_cv(record.profile)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"CV analysis failed: {exc}")
    # Structured path: set intelligence, clear any stale raw text
    # Fallback path: set raw text, clear any stale structured intelligence
    record.cv_intelligence = result.intelligence
    record.cv_intelligence_raw = result.raw_text if result.parse_failed else None
    record.cv_intelligence_generated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.post("/clients/{client_id}/generate-positioning")
def post_generate_positioning(client_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    try:
        result = generate_positioning(record.profile, record.cv_intelligence)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Positioning generation failed: {exc}")
    record.positioning = result.positioning
    record.positioning_raw = result.raw_text if result.parse_failed else None
    record.positioning_generated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.post("/clients/{client_id}/run-market-radar")
def post_run_market_radar(client_id: str, body: MarketRadarRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    try:
        result = run_market_radar(
            record.profile,
            cv_intelligence=record.cv_intelligence,
            cv_intelligence_raw=record.cv_intelligence_raw,
            positioning=record.positioning,
            positioning_raw=record.positioning_raw,
            manual_research=body.manual_research,
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Market Radar failed: {exc}")
    if result.is_complete and not result.parse_failed:
        # Complete structured result — save everything, clear any previous warning
        record.market_radar = result.radar
        record.market_radar_raw = None
        record.market_radar_generated_at = datetime.now(timezone.utc).isoformat()
        record.market_radar_is_complete = True
        record.market_radar_scan_warning = None
    elif record.market_radar_is_complete:
        # Incomplete result but a previous complete scan exists — preserve it, surface warning only
        sections = ", ".join(result.missing_sections) if result.missing_sections else "multiple sections"
        record.market_radar_scan_warning = (
            f"Latest scan returned incomplete results (missing: {sections}). "
            "Showing previous complete scan."
        )
    else:
        # Incomplete result, no previous complete scan — save as draft
        record.market_radar = result.radar
        record.market_radar_raw = result.raw_text if result.parse_failed else None
        record.market_radar_generated_at = datetime.now(timezone.utc).isoformat()
        record.market_radar_is_complete = False
        if result.missing_sections:
            sections = ", ".join(result.missing_sections)
            record.market_radar_scan_warning = (
                f"Scan returned partial results — some sections are incomplete ({sections}). "
                "Try refreshing the scan."
            )
        else:
            record.market_radar_scan_warning = (
                "Scan returned partial results. Try refreshing the scan."
            )
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


# ── Opportunity endpoints ─────────────────────────────────────────────────────


@app.post("/clients/{client_id}/opportunities")
def post_create_opportunity(client_id: str, body: OpportunityRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    opp = Opportunity(
        title=body.title,
        company=body.company,
        pathway=body.pathway,
        source_type=body.source_type,
        source_section=body.source_section,
        confidence=body.confidence,
        priority=body.priority,
        status=body.status,
        fit_rationale=body.fit_rationale,
        evidence=body.evidence,
        relationship_route=body.relationship_route,
        next_action=body.next_action,
        advisor_note=body.advisor_note,
        sources=body.sources,
    )
    record.opportunities.append(opp)
    updated = update_client(record)
    return JSONResponse(status_code=201, content=updated.model_dump(mode="json"))


@app.put("/clients/{client_id}/opportunities/{opp_id}")
def put_opportunity(client_id: str, opp_id: str, body: OpportunityRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    opp = next((o for o in record.opportunities if o.id == opp_id), None)
    if opp is None:
        raise HTTPException(status_code=404, detail="Opportunity not found.")
    opp.title = body.title
    opp.company = body.company
    opp.pathway = body.pathway
    opp.source_type = body.source_type
    opp.source_section = body.source_section
    opp.confidence = body.confidence
    opp.priority = body.priority
    opp.status = body.status
    opp.fit_rationale = body.fit_rationale
    opp.evidence = body.evidence
    opp.relationship_route = body.relationship_route
    opp.next_action = body.next_action
    opp.advisor_note = body.advisor_note
    opp.sources = body.sources
    opp.updated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.delete("/clients/{client_id}/opportunities/{opp_id}")
def delete_opportunity(client_id: str, opp_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    before = len(record.opportunities)
    record.opportunities = [o for o in record.opportunities if o.id != opp_id]
    if len(record.opportunities) == before:
        raise HTTPException(status_code=404, detail="Opportunity not found.")
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


# ── Session Note endpoints ───────────────────────────────────────────────────


@app.post("/clients/{client_id}/notes")
def post_create_note(client_id: str, body: SessionNoteRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    note = SessionNote(
        date=body.date, title=body.title,
        notes=body.notes, advisor_only=body.advisor_only,
    )
    record.session_notes.append(note)
    updated = update_client(record)
    return JSONResponse(status_code=201, content=updated.model_dump(mode="json"))


@app.put("/clients/{client_id}/notes/{note_id}")
def put_note(client_id: str, note_id: str, body: SessionNoteRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    note = next((n for n in record.session_notes if n.id == note_id), None)
    if note is None:
        raise HTTPException(status_code=404, detail="Note not found.")
    note.date = body.date
    note.title = body.title
    note.notes = body.notes
    note.advisor_only = body.advisor_only
    note.updated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.delete("/clients/{client_id}/notes/{note_id}")
def delete_note(client_id: str, note_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    before = len(record.session_notes)
    record.session_notes = [n for n in record.session_notes if n.id != note_id]
    if len(record.session_notes) == before:
        raise HTTPException(status_code=404, detail="Note not found.")
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


# ── Action Item endpoints ─────────────────────────────────────────────────────


@app.post("/clients/{client_id}/actions")
def post_create_action(client_id: str, body: ActionItemRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    action = ActionItem(
        action=body.action, owner=body.owner,
        due_date=body.due_date, status=body.status,
        related_opportunity=body.related_opportunity,
        advisor_note=body.advisor_note,
    )
    record.action_items.append(action)
    updated = update_client(record)
    return JSONResponse(status_code=201, content=updated.model_dump(mode="json"))


@app.put("/clients/{client_id}/actions/{action_id}")
def put_action(client_id: str, action_id: str, body: ActionItemRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    action = next((a for a in record.action_items if a.id == action_id), None)
    if action is None:
        raise HTTPException(status_code=404, detail="Action item not found.")
    action.action = body.action
    action.owner = body.owner
    action.due_date = body.due_date
    action.status = body.status
    action.related_opportunity = body.related_opportunity
    action.advisor_note = body.advisor_note
    action.updated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.delete("/clients/{client_id}/actions/{action_id}")
def delete_action(client_id: str, action_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    before = len(record.action_items)
    record.action_items = [a for a in record.action_items if a.id != action_id]
    if len(record.action_items) == before:
        raise HTTPException(status_code=404, detail="Action item not found.")
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


# ── Target Contact endpoints ──────────────────────────────────────────────────


@app.post("/clients/{client_id}/target-contacts/search")
def post_search_contacts(client_id: str, body: ContactSearchRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    try:
        result = search_contacts(
            company=body.company,
            role_context=body.role_context,
            search_focus=body.search_focus,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Contact search failed: {exc}")
    return JSONResponse(status_code=200, content={
        "contacts": [
            {
                "name": c.name, "title": c.title, "company": c.company,
                "linkedin_url": c.linkedin_url, "source_url": c.source_url,
                "why_relevant": c.why_relevant, "suggested_angle": c.suggested_angle,
                "confidence": c.confidence,
            }
            for c in result.contacts
        ],
        "search_mode": result.search_mode,
        "message": result.message,
    })


@app.post("/clients/{client_id}/target-contacts/extract-from-notes")
def post_extract_from_notes(client_id: str, body: ExtractFromNotesRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    if not body.notes.strip():
        raise HTTPException(status_code=422, detail="Notes are required.")
    try:
        result = extract_network_from_notes(record, body.notes)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Extraction failed: {exc}")
    return JSONResponse(status_code=200, content={
        "suggested_contacts": [
            {
                "name": c.name,
                "current_title": c.current_title,
                "company": c.company,
                "network_source": c.network_source,
                "relationship_owner": c.relationship_owner,
                "relationship_to_client": c.relationship_to_client,
                "relationship_to_advisor": c.relationship_to_advisor,
                "relationship_strength": c.relationship_strength,
                "role_in_search": c.role_in_search,
                "target_company": c.target_company,
                "target_sector": c.target_sector,
                "bridge_to": c.bridge_to,
                "warm_path_status": c.warm_path_status,
                "ask_type": c.ask_type,
                "suggested_approach": c.suggested_approach,
                "opportunity_path_hypothesis": c.opportunity_path_hypothesis,
                "next_action": c.next_action,
                "next_action_owner": c.next_action_owner,
                "status": c.status,
                "advisor_only": c.advisor_only,
                "client_shareable": c.client_shareable,
                "approved_for_outreach": c.approved_for_outreach,
                "include_in_advisor_brief": c.include_in_advisor_brief,
                "include_in_weekly_plan": c.include_in_weekly_plan,
                "missing_information": c.missing_information,
                "follow_up_questions": c.follow_up_questions,
                "confidence": c.confidence,
                "evidence_from_notes": c.evidence_from_notes,
            }
            for c in result.suggested_contacts
        ],
        "network_insights": result.network_insights,
        "recommended_follow_up_questions": result.recommended_follow_up_questions,
    })


_TARGET_CONTACT_REQUEST_FIELDS = list(TargetContactRequest.model_fields.keys())


@app.post("/clients/{client_id}/target-contacts")
def post_create_contact(client_id: str, body: TargetContactRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    contact = TargetContact(**{f: getattr(body, f) for f in _TARGET_CONTACT_REQUEST_FIELDS})
    record.target_contacts.append(contact)
    updated = update_client(record)
    return JSONResponse(status_code=201, content=updated.model_dump(mode="json"))


@app.put("/clients/{client_id}/target-contacts/{contact_id}")
def put_contact(client_id: str, contact_id: str, body: TargetContactRequest):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    contact = next((c for c in record.target_contacts if c.id == contact_id), None)
    if contact is None:
        raise HTTPException(status_code=404, detail="Contact not found.")
    for f in _TARGET_CONTACT_REQUEST_FIELDS:
        setattr(contact, f, getattr(body, f))
    contact.updated_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.delete("/clients/{client_id}/target-contacts/{contact_id}")
def delete_contact(client_id: str, contact_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    before = len(record.target_contacts)
    record.target_contacts = [c for c in record.target_contacts if c.id != contact_id]
    if len(record.target_contacts) == before:
        raise HTTPException(status_code=404, detail="Contact not found.")
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


# ── Advisor Brief endpoint ────────────────────────────────────────────────────


@app.post("/clients/{client_id}/generate-advisor-brief")
def post_generate_advisor_brief(client_id: str):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    try:
        result = generate_advisor_brief(record)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Advisor Brief generation failed: {exc}")
    record.advisor_brief = result.brief
    record.advisor_brief_raw = result.raw_text if result.parse_failed else None
    record.advisor_brief_generated_at = datetime.now(timezone.utc).isoformat()
    record.advisor_brief_is_edited = False
    record.advisor_brief_edited_at = None
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))


@app.put("/clients/{client_id}/advisor-brief")
def put_advisor_brief(client_id: str, body: AdvisorBrief):
    record = get_client(client_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Client not found.")
    record.advisor_brief = body
    record.advisor_brief_is_edited = True
    record.advisor_brief_edited_at = datetime.now(timezone.utc).isoformat()
    updated = update_client(record)
    return JSONResponse(status_code=200, content=updated.model_dump(mode="json"))